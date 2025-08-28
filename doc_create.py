from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
import re


# 清理文本中的非法Unicode字符，特别是代理对字符
def clean_text_for_utf8(text):
    if text is None:
        return ""

    try:
        # 尝试编码为utf-8，如果失败则替换问题字符
        cleaned_text = text.encode('utf-8', errors='replace').decode('utf-8')

        # 额外清理可能导致问题的Unicode代理对字符
        # 代理对范围是 U+D800 到 U+DFFF
        cleaned_text = re.sub(r'[\uD800-\uDFFF]', '?', cleaned_text)

        return cleaned_text
    except Exception as e:
        print(f"清理文本时出错: {e}")
        # 如果出现任何错误，返回一个安全的字符串
        return "[文本包含无法显示的字符]"


# for word convert
def set_font(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_title(doc, text, level=1, bold=True, font_name='微软雅黑', font_size=12):
    # 清理文本
    cleaned_text = clean_text_for_utf8(text)

    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(cleaned_text)
    run.bold = bold
    set_font(run, font_name, font_size)
    return paragraph


def add_paragraph(doc, text, font_name='微软雅黑', font_size=12, bold=False):
    # 清理文本
    cleaned_text = clean_text_for_utf8(text)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(cleaned_text)
    set_font(run, font_name, font_size)
    run.bold = bold
    return paragraph


def add_link(doc, text, url, font_name='微软雅黑', font_size=12):
    # 清理文本
    cleaned_text = clean_text_for_utf8(text)

    paragraph = doc.add_paragraph()
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                          is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run for the hyperlink
    run = paragraph.add_run(cleaned_text)
    run.font.color.rgb = OxmlElement('w:color').set(qn('w:val'), '0000FF')
    run.font.underline = True
    set_font(run, font_name, font_size)

    hyperlink.append(run._r)
    paragraph._element.append(hyperlink)

    return paragraph

#
# def add_bullet_points(doc, items, font_name='微软雅黑', font_size=12):
#     # 遍历要添加的项目符号列表
#     for item in items:
#         # 添加一个新段落
#         paragraph = doc.add_paragraph()
#         # 设置段落的样式为项目符号
#         paragraph.style = 'List Bullet'
#         # 添加文本到段落
#         run = paragraph.add_run(item)
#         # 设置字体和字号
#         run.font.name = font_name
#         run.font.size = Pt(font_size)
