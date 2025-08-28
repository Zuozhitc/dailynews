import functions as func
import doc_create as doc
from docx import Document
import catch_news
import datetime
import arxiv_summary as arxiv_func
from catch_sources import catch_ProductHunt
import os

# 可选：用于 fallback 的简单 docx 读取
try:
    import mammoth
    HAS_MAMMOTH = True
except Exception:
    HAS_MAMMOTH = False


# 保存文件word，避免文件被占用的时候报错，保存文件名字后加日期时间
# ✅ 现在会返回保存的文件路径
def save_doc(new_doc, phase_text=''):
    year = datetime.datetime.now().year
    now = datetime.datetime.now()
    # 保存在当前目录下 AI_Report文件夹内。如果文件夹不存在，则创建文件夹
    path = f'AI_Report{year}'
    if not os.path.exists(path):
        os.makedirs(path)
    # 为当天创建子目录
    path = path + now.strftime("/%m%d")
    if not os.path.exists(path):
        os.makedirs(path)

    filename = os.path.join(path, f"AI_daily_report_{phase_text}{now.strftime('%m%d-%H')}.docx")
    try:
        new_doc.save(filename)
    except PermissionError:
        filename = os.path.join(path, f"AI_daily_report_{phase_text}{now.strftime('%m%d-%H')}_temp.docx")
        new_doc.save(filename)
    return filename


# docx → html
def convert_docx_to_html(docx_path) -> str:
    """
    首选用 mammoth 转换，高保真 HTML
    如果 mammoth 不可用，fallback 为极简段落拼接（保格式能力有限）
    """
    if HAS_MAMMOTH:
        try:
            with open(docx_path, "rb") as f:
                result = mammoth.convert_to_html(f)
                return result.value or ""
        except Exception as e:
            print(f"mammoth 转换出错：{e}")

    # fallback：极简，把段落转 <p>（不含链接样式、标题样式等）
    try:
        from docx import Document as _Doc
        d = _Doc(docx_path)
        parts = []
        for p in d.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(f"<p>{text}</p>")
        return "\n".join(parts)
    except Exception as e:
        print(f"fallback 转换出错：{e}")
        return ""


# 分批粗筛选标题
def rough_filter_batch(new_infos) -> list:
    # 把titles写到一个str中
    info_titles = [f"{index}. {item['title']}" for index, item in enumerate(new_infos)]
    info_titles_text = '\n'.join(info_titles)

    # 让gpt-5筛选哪些titles是相关内容，输出索引号
    filtered_info_titles_index = func.title_filter_rough(info_titles_text)

    # 筛选出包含这些有用title的元素
    filtered_info = [item for index, item in enumerate(new_infos) if index in filtered_info_titles_index]
    return filtered_info


# 获取各大新闻网站的新闻信息，并进行筛选，返回筛选后的新闻标题和链接的对应关系
def get_news(days=1):
    print('STEP 1 开启爬虫，获取最近的新闻')
    print('STEP 2 开始获取最新新闻，目前默认为24小时左右的新闻')
    new_infos = catch_news.catch_all(days)
    print('^_^获取本周新闻成功，去重后，总共' + str(len(new_infos)) + '条\n\n')

    # 第一次筛选
    print('STEP 3 第一次筛选新闻标题')
    filtered_info = []
    batch_size = 30
    for i in range(0, len(new_infos), batch_size):
        batch = new_infos[i: i + batch_size]
        filtered_info.extend(rough_filter_batch(batch))
    print('^_^第一次筛选成功，一共有' + str(len(filtered_info)) + '条新闻\n\n')

    for info in filtered_info:
        print(info['title'], info['link'])

    # 第二次筛选与分组
    print('STEP 4 第二次筛选新闻标题')
    info_titles2link = {item['title']: item['link'] for item in filtered_info}
    filtered_info_index_titles = [f"{index}. {item['title']}" for index, item in enumerate(filtered_info)]
    filtered_info_index_titles_text = '\n'.join(filtered_info_index_titles)

    info_titles_final_index = func.title_filter_final(filtered_info_index_titles_text)
    print("info_titles_final_index:", info_titles_final_index)

    # 用内容替换序号，得到最终的标题结构
    info_titles_final = {}
    for category, indices_dict in info_titles_final_index.items():
        if isinstance(indices_dict, dict):
            cleaned_indices_dict = {}
            for subcategory, indices in indices_dict.items():
                unique_titles = list({filtered_info[i]['title']: i for i in indices}.keys())
                if unique_titles:
                    cleaned_indices_dict[subcategory] = {title: i for i, title in enumerate(unique_titles)}
            if cleaned_indices_dict:
                info_titles_final[category] = cleaned_indices_dict

    print('^_^处理新闻资讯成功，新闻架构如下')
    print("info_titles_final:", info_titles_final)
    return info_titles_final, info_titles2link


# 获取论文信息
def get_papers(paper_days=1, save_to_excel=False):
    print('STEP 5 开始获取论文信息')
    try:
        papers = arxiv_func.get_and_write_excel(paper_days, save_to_excel=save_to_excel)
        if papers is None:
            papers = []
        print('^_^获取API论文信息成功，总共' + str(len(papers)) + '篇论文\n\n')

        # 如果API没有返回论文或者第一篇已存在，尝试从网页获取
        if len(papers) == 0 or (len(papers) > 0 and catch_news.link_exists(papers[0])):
            print('From API: No new papers found in the last', paper_days, 'days.')
            try:
                import arxiv_summary_web
                web_papers = arxiv_summary_web.from_web_to_list_advanced()
                if web_papers is not None and len(web_papers) > 0:
                    papers = web_papers
                    print('^_^获取WEB论文信息成功，总共' + str(len(papers)) + '篇论文\n\n')
                else:
                    print('从网页获取论文也未找到新内容')
                    papers = []
            except Exception as e:
                print(f"从网页获取论文时出错: {e}")
                papers = []

        if papers is None:
            papers = []

        # 去重
        papers = catch_news.get_unique_links(papers)
        return papers
    except Exception as e:
        print(f"获取论文信息时出错: {e}")
        return []


# 获取ProductHunt信息
def get_producthunt():
    print('STEP 6 开始获取ProductHunt信息')
    try:
        # 获取完整的产品信息列表（包含title、link、tag等）
        producthunt_items = catch_ProductHunt.catch_producthunt()

        if producthunt_items is None:
            producthunt_items = []

        # 提取链接用于去重检查
        producthunt_links = [item['link'] for item in producthunt_items]
        unique_links = catch_news.get_unique_links(producthunt_links)

        # 只保留链接未出现过的产品
        filtered_producthunt = []
        for item in producthunt_items:
            if item['link'] in unique_links:
                filtered_producthunt.append(item)

        print('^_^获取ProductHunt信息成功，总共' + str(len(filtered_producthunt)) + '条AI相关产品信息\n\n')
        return filtered_producthunt
    except Exception as e:
        print(f"获取ProductHunt信息时出错: {e}")
        return []


# 用于将新闻玩家信息总结
def write_news(new_doc, info_titles_final, info_titles2link, manual=False, choice='gpt5'):
    # 标准开头
    doc.add_title(new_doc, f'【ChatGPT全球动态日报{datetime.datetime.now().strftime("%m%d")}】 — TEG战略发展中心推送', level=1)
    doc.add_paragraph(new_doc, 'Dear all，')
    doc.add_paragraph(new_doc, '该报是由战略发展中心推送，帮助大家同步全球范围内 ChatGPT 新闻动态。报告内容分为五大部分，其中包含非公开信息，仅在团队内部分享：')
    doc.add_paragraph(new_doc, '一、 玩家动态追踪', bold=True)
    doc.add_paragraph(new_doc, '二、 技术前沿分析', bold=True)
    doc.add_paragraph(new_doc, '三、 应用场景分析', bold=True)
    doc.add_paragraph(new_doc, '四、 内部论坛', bold=True)
    doc.add_paragraph(new_doc, '五、 论文及非公开信息', bold=True)
    doc.add_paragraph(new_doc, '大家如果有任何问题和建议，请联系 TEG 战略发展中心。')
    doc.add_paragraph(new_doc, '以下是正文：')

    # 新增去重能力。避免同一个链接被反复写入
    link_set = set()

    # 开始写入内容
    doc.add_title(new_doc, '一、玩家动态追踪', level=1)
    if info_titles_final.get('重点玩家'):
        for player in list(info_titles_final['重点玩家'].keys()):
            doc.add_title(new_doc, player, level=2)

            for player_info in info_titles_final['重点玩家'][player]:
                # 去重
                if info_titles2link[player_info] in link_set:
                    continue
                else:
                    link_set.add(info_titles2link[player_info])
                    catch_news.add_links([info_titles2link[player_info]])  # 将链接加入到数据库中，避免重复
                spaced_player_info = func.text_format_beautify(player_info)
                print(spaced_player_info)
                summary = func.summarize_web(info_titles2link[player_info], choice)

                # 标题用原文标题 + 链接
                doc.add_link(new_doc, spaced_player_info, info_titles2link[player_info])
                if manual:
                    # 便于人工发送日报时复制
                    doc.add_paragraph(new_doc, info_titles2link[player_info])
                    doc.add_paragraph(new_doc, "")

                # 摘要正文
                for line in summary.split('\n'):
                    doc.add_paragraph(new_doc, line)
                doc.add_paragraph(new_doc, '')
    else:
        doc.add_paragraph(new_doc, '无')

    # 分阶段保存：news
    save_doc(new_doc, 'news')


# 用于将论文信息总结
def write_papers(new_doc, papers, manual=False, choice='gpt5'):
    doc.add_title(new_doc, '二、技术前沿分析', level=1)

    # 无论文
    if papers is None or len(papers) == 0:
        doc.add_paragraph(new_doc, '无新论文')
        save_doc(new_doc, 'papers')
        return

    paper_count = 0
    for paper in papers:
        try:
            if not catch_news.link_exists(paper):
                summary = func.summarize_paper(paper, choice)
                catch_news.add_links([paper])

                if summary is None or len(summary.strip()) == 0:
                    continue

                # 分割标题和内容
                summary_parts = summary.split('\n')
                title = summary_parts[0]
                summary_content = summary_parts[1:] if len(summary_parts) > 1 else []

                doc.add_link(new_doc, title, paper)
                if manual:
                    doc.add_paragraph(new_doc, paper)
                    doc.add_paragraph(new_doc, "")

                for line in summary_content:
                    doc.add_paragraph(new_doc, line)
                doc.add_paragraph(new_doc, '')
                paper_count += 1
        except Exception as e:
            print(f"处理论文时出错: {e}")
            continue

    if paper_count == 0:
        doc.add_paragraph(new_doc, '无新论文')

    # 分阶段保存：papers
    save_doc(new_doc, 'papers')


def write_producthunt(new_doc, producthunt_items, manual=False, choice='gpt5'):
    doc.add_title(new_doc, '三、应用场景分析', level=1)

    if producthunt_items is None or len(producthunt_items) == 0:
        doc.add_paragraph(new_doc, '无新AI产品')
        save_doc(new_doc, 'product')
        return

    product_count = 0
    for item in producthunt_items:
        try:
            link = item['link']
            original_title = item.get('title', '')

            if not catch_news.link_exists(link):
                catch_news.add_links([link])
                summary = func.summarize_product(link, choice)

                if summary is None or len(summary.strip()) == 0:
                    continue

                lines = summary.strip().split('\n')
                display_title = lines[0].strip()  # 例如 "Predict，AI 模型评估平台"
                content_lines = lines[1:] if len(lines) > 1 else []

                # 展示标题（小节）
                doc.add_paragraph(new_doc, display_title, bold=True, font_size=11)

                # 保留原文链接、PH 标题（便于人工核对）
                if manual:
                    doc.add_link(new_doc, "原文链接", link)
                    if original_title:
                        doc.add_paragraph(new_doc, f"ProductHunt 标题: {original_title}")
                    doc.add_paragraph(new_doc, "")

                # 正文内容
                for line in content_lines:
                    if line.strip():
                        doc.add_paragraph(new_doc, line)

                doc.add_paragraph(new_doc, '')
                product_count += 1
        except Exception as e:
            print(f"处理产品时出错: {e}")
            continue

    if product_count == 0:
        doc.add_paragraph(new_doc, '无新AI产品')

    # 分阶段保存：product
    save_doc(new_doc, 'product')


def make_report(manual=False, paper_days=1, news_days=1,
                makenews=True, makepapers=True, makeproduct=True, choice='gpt5'):
    """
    生成报告（docx），并在最后把最终 docx 转为 HTML 字符串返回
    scheduler → send_email 会把该 HTML 作为邮件正文发送
    """
    try:
        # 初始化word文档
        new_doc = Document()

        # Part 1: 新闻
        if makenews:
            try:
                info_titles_final, info_titles2link = get_news(news_days)
                write_news(new_doc, info_titles_final, info_titles2link, manual, choice=choice)
            except Exception as e:
                print(f"处理新闻部分时出错: {e}")
                doc.add_title(new_doc, '一、玩家动态追踪', level=1)
                doc.add_paragraph(new_doc, '处理新闻时出错，请检查日志')

        # Part 2: 论文
        if makepapers:
            try:
                papers = get_papers(paper_days, save_to_excel=True)
                write_papers(new_doc, papers, manual, choice=choice)
            except Exception as e:
                print(f"处理论文部分时出错: {e}")
                doc.add_title(new_doc, '二、技术前沿分析', level=1)
                doc.add_paragraph(new_doc, '处理论文时出错，请检查日志')

        # Part 3: ProductHunt
        if makeproduct:
            try:
                producthunt_items = get_producthunt()
                write_producthunt(new_doc, producthunt_items, manual, choice=choice)
            except Exception as e:
                print(f"处理产品部分时出错: {e}")
                doc.add_title(new_doc, '三、应用场景分析', level=1)
                doc.add_paragraph(new_doc, '处理产品时出错，请检查日志')

        # ✅ 不论是否全开，都保存一份 final，并将其转 HTML
        final_path = save_doc(new_doc, 'final')
        html_content = convert_docx_to_html(final_path)

        print('All Done!')
        return html_content if html_content is not None else ""
    except Exception as e:
        print(f"生成报告时出错: {e}")
        return ""


# 用于辅助人工发送日报，复制粘贴链接。默认为False（走 scheduler 时一般为 False）
manual = True
# 默认获取1天的论文。因为arxiv更新延迟，如果今天是周二，则获取3天的论文
paper_days = 1 if datetime.datetime.now().weekday() != 1 else 3
# 默认获取24小时内的新闻，如果今天是周一，则获取3天的新闻（有去重，所以多获取一天）
news_days = 2 if datetime.datetime.now().weekday() != 0 else 3

# make_report(manual, paper_days, news_days)  # 调试时可手动运行
